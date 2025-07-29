import os
from typing import Dict, List, Optional

import chromadb
import PyPDF2
import streamlit as st
from docx import Document as DocxDocument
from llama_index.core import Document, Settings, StorageContext, VectorStoreIndex
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore


class DocumentManager:
    def __init__(self, data_dir: str = "data"):
        self.data_dir = data_dir
        self.documents_dir = os.path.join(data_dir, "documents")
        self.index_dir = os.path.join(data_dir, "index")

        # 创建目录
        os.makedirs(self.documents_dir, exist_ok=True)
        os.makedirs(self.index_dir, exist_ok=True)

        self._initialize_index()

    @st.cache_resource
    def _initialize_index(_self):
        """初始化索引（使用Streamlit缓存）"""
        try:
            # 初始化ChromaDB
            chroma_client = chromadb.PersistentClient(path=_self.index_dir)
            chroma_collection = chroma_client.get_or_create_collection("documents")

            # 配置embedding模型
            embed_model = HuggingFaceEmbedding(
                model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
            )

            # 配置服务上下文
            # service_context = ServiceContext.from_defaults(
            #     embed_model=embed_model,
            #     chunk_size=512,
            #     chunk_overlap=50
            # )
            Settings.embed_model = embed_model
            Settings.chunk_size = 512
            Settings.chunk_overlap = 50

            # 配置向量存储
            vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
            storage_context = StorageContext.from_defaults(vector_store=vector_store)

            index = VectorStoreIndex.from_documents(
                [],
                embed_model=embed_model,
                storage_context=storage_context
            )

            return index, embed_model
        except Exception as e:
            st.error(f"初始化索引失败: {e}")
            return None, None

    def _extract_text_from_file(self, file_path: str) -> str:
        """从文件中提取文本"""
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext in ['.txt', '.md']:
                with open(file_path, encoding='utf-8') as f:
                    return f.read()

            elif file_ext == '.pdf':
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text

            elif file_ext == '.docx':
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text

            else:
                return f"不支持的文件格式: {file_ext}"

        except Exception as e:
            return f"文件读取错误: {str(e)}"

    def add_document(self, file_path: str, category: str = "general") -> bool:
        """添加文档到索引"""
        try:
            index, service_context = self._initialize_index()
            if not index:
                return False

            # 支持的文件类型
            supported_extensions = ['.txt', '.md', '.pdf', '.docx']
            file_ext = os.path.splitext(file_path)[1].lower()

            if file_ext not in supported_extensions:
                st.error(f"不支持的文件格式: {file_ext}")
                return False

            # 提取文档内容
            content = self._extract_text_from_file(file_path)

            if not content.strip():
                st.error("文档内容为空")
                return False

            # 创建文档对象
            document = Document(
                text=content,
                metadata={
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "category": category,
                    "file_type": file_ext,
                    "file_size": os.path.getsize(file_path)
                }
            )

            # 添加到索引
            index.insert(document)

            # 保存文档副本
            import shutil
            dest_path = os.path.join(self.documents_dir, os.path.basename(file_path))
            if file_path != dest_path:
                shutil.copy2(file_path, dest_path)

            return True

        except Exception as e:
            st.error(f"添加文档失败: {e}")
            return False

    def search_documents(self, query: str, category: Optional[str] = None, top_k: int = 5) -> List[Dict]:
        """搜索文档"""
        try:
            index, service_context = self._initialize_index()
            if not index:
                return []

            query_engine = index.as_query_engine(
                similarity_top_k=top_k,
                service_context=service_context
            )

            response = query_engine.query(query)

            results = []
            for node in response.source_nodes:
                metadata = node.metadata
                result = {
                    "content": node.text[:300] + "..." if len(node.text) > 300 else node.text,
                    "file_name": metadata.get("file_name", ""),
                    "category": metadata.get("category", ""),
                    "file_type": metadata.get("file_type", ""),
                    "file_size": metadata.get("file_size", 0),
                    "score": getattr(node, 'score', 0)
                }

                # 分类过滤
                if category and category != "全部" and result["category"] != category:
                    continue

                results.append(result)

            return results

        except Exception as e:
            st.error(f"搜索文档失败: {e}")
            return []

    def get_categories(self) -> List[str]:
        """获取所有分类"""
        return ["general", "work", "study", "personal", "research"]

    def get_document_stats(self) -> Dict:
        """获取文档统计信息"""
        stats = {
            "total_files": 0,
            "categories": {},
            "file_types": {},
            "total_size": 0
        }

        if os.path.exists(self.documents_dir):
            for file_name in os.listdir(self.documents_dir):
                file_path = os.path.join(self.documents_dir, file_name)
                if os.path.isfile(file_path):
                    stats["total_files"] += 1
                    stats["total_size"] += os.path.getsize(file_path)

                    # 文件类型统计
                    ext = os.path.splitext(file_name)[1].lower()
                    stats["file_types"][ext] = stats["file_types"].get(ext, 0) + 1

        return stats
